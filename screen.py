from mss import mss
from docx import Document
from docx.shared import Inches
import os
from datetime import datetime

def append_screenshot_to_docx(docx_name="report.docx"):
    temp_img = "temp_shot.png"
    
    # 1. Делаем скриншот первого монитора
    with mss() as sct:
        sct.shot(mon=1, output=temp_img)

    # 2. Проверяем наличие файла: открываем существующий или создаем новый
    if os.path.exists(docx_name):
        doc = Document(docx_name)
    else:
        doc = Document()
        doc.add_heading('Журнал скриншотов', 0)

    # 3. Добавляем временную метку
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    doc.add_paragraph(f"Снимок сделан: {timestamp}")

    # 4. Вставляем изображение
    # Если скриншоты слишком большие, можно уменьшить Inches(6.0)
    doc.add_picture(temp_img, width=Inches(6.0))

    # 5. Сохраняем документ
    doc.save(docx_name)
    
    # Удаляем временный файл картинки
    if os.path.exists(temp_img):
        os.remove(temp_img)
        
    print(f"Скриншот успешно добавлен в {docx_name}")

# Запуск функции
append_screenshot_to_docx()
