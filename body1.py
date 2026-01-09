import json
import time
import pyautogui
import requests
import subprocess
import webbrowser
import re
import sys
import os
from datetime import datetime
from mss import mss
from docx import Document
from docx.shared import Inches

# --- КОНФИГУРАЦИЯ API ---
API_KEY = "ВАШ_API_KEY" # Не забудьте вставить ваш ключ
BASE_URL = "https://openrouter.ai/api/v1/chat/completions"
MODEL = "kwaipilot/kat-coder-pro:free"

# --- ЛОГИКА ОТЧЕТОВ (бывший screen.py) ---
def append_screenshot_to_docx(docx_name="report.docx"):
    """Делает скриншот и сохраняет его в Word-документ."""
    temp_img = "temp_shot.png"
    try:
        # 1. Делаем скриншот первого монитора
        with mss() as sct:
            sct.shot(mon=1, output=temp_img)

        # 2. Проверяем наличие файла: открываем существующий или создаем новый
        if os.path.exists(docx_name):
            doc = Document(docx_name)
        else:
            doc = Document()
            doc.add_heading('Журнал скриншотов', 0)

        # 3. Добавляем временную метку
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        doc.add_paragraph(f"Снимок сделан: {timestamp}")

        # 4. Вставляем изображение
        doc.add_picture(temp_img, width=Inches(6.0))

        # 5. Сохраняем документ
        doc.save(docx_name)
        
        # Удаляем временный файл картинки
        if os.path.exists(temp_img):
            os.remove(temp_img)
        print(f"[+] Скриншот успешно добавлен в {docx_name}")
    except Exception as e:
        print(f"[-] Ошибка при создании скриншота: {e}")

# --- ОБНОВЛЕННЫЙ ПРОМПТ ---
SYSTEM_PROMPT = """
Ты — модуль управления Red OS. Возвращай ТОЛЬКО валидный JSON-список команд БЕЗ текста и markdown.

Формат: [{"action": "hotkey", "params": ["ctrl","alt","t"]}, ...]

**Действия:**
- "hotkey": ["key1", "key2"] (нажать клавиши)
- "type": ["текст"] (ввод текста)
- "wait": [секунды] (пауза)
- "run_shell": ["команда"] (выполнить в фоне)
- "open_url": ["ссылка"] (открыть браузер)
- "mouse_click": [x, y] (клик мышью)
- "screenshot": [] (сделать скриншот экрана и сохранить в отчет docx)

**Правила Red OS:**
- Терминал: ctrl+alt+t → type → enter
- Пакеты: dnf вместо apt
- Root: su - → enter → type "1" → enter
- Рабочий стол: /home/sergeykos/Desktop
- Минимум 1 команда, максимум 10
- Воспринимай пункты списка (1. 2.) как отдельные инструкции
- После каждой интрукции запускай логику "screenshot"
"""

def get_ai_instruction(user_text: str, max_retries: int = 4) -> list | None:
    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": MODEL,
        "messages": [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": f"Задача: {user_text}"},
        ],
        "temperature": 0.1,
    }

    for attempt in range(max_retries):
        try:
            resp = requests.post(BASE_URL, headers=headers, json=payload, timeout=30)
            if resp.status_code == 429:
                time.sleep(2 ** attempt)
                continue
            
            resp.raise_for_status()
            content = resp.json()["choices"][0]["message"]["content"].strip()
            content = re.sub(r'```json\s*|```|\`{1,3}', '', content).strip()
            parsed = json.loads(content)
            if isinstance(parsed, list):
                return parsed
        except Exception as e:
            print(f"[-] Ошибка AI (попытка {attempt+1}): {e}")
            time.sleep(1)
    return None

def execute_command_list(command_list: list):
    print(f"[*] Выполняю {len(command_list)} команд:")
    for i, cmd in enumerate(command_list, 1):
        action = cmd.get("action")
        params = cmd.get("params", [])
        print(f"  Шаг {i}: {action} {params}")
        
        try:
            if action == "hotkey":
                pyautogui.hotkey(*params)
            elif action == "type":
                if params:
                    pyautogui.write(str(params[0]), interval=0.015)
            elif action == "wait":
                wait_time = float(params[0]) if params else 1.0
                time.sleep(wait_time)
            elif action == "run_shell":
                if params:
                    subprocess.Popen(str(params[0]), shell=True, 
                                   stdout=subprocess.DEVNULL, 
                                   stderr=subprocess.DEVNULL)
            elif action == "open_url":
                if params:
                    webbrowser.open(str(params[0]))
            elif action == "mouse_click" and len(params) == 2:
                pyautogui.click(int(params[0]), int(params[1]))
            elif action == "screenshot":
                append_screenshot_to_docx()
        except Exception as e:
            print(f"    ❌ Ошибка шага {i}: {e}")

def main():
    pyautogui.FAILSAFE = True
    pyautogui.PAUSE = 0.1
    print("🤖 Бот Red OS (Интегрированная версия)")
    print("Доступно новое действие: 'screenshot'")
    
    while True:
        try:
            user_input = input("\n> Задача: ").strip()
            if user_input.lower() in ["exit", "quit", "выход"]:
                break
            if not user_input:
                continue

            tasks = re.split(r'(\d+\.\s+)', user_input)
            tasks = [t.strip() for t in tasks if t.strip() and not re.match(r'^\d+\.$', t.strip())]

            for idx, task in enumerate(tasks, 1):
                print(f"\n📋 Подзадача {idx}/{len(tasks)}: {task}")
                commands = get_ai_instruction(task)
                if commands:
                    execute_command_list(commands)
                else:
                    print(f"[-] Не удалось получить инструкции для: {task}")

        except KeyboardInterrupt:
            break
        except Exception as e:
            print(f"❌ Критическая ошибка: {e}")

if __name__ == "__main__":
    main()
